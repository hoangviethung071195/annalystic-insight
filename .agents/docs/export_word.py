import os
import subprocess
import sys

# Tự động cài đặt thư viện python-docx nếu chưa có
try:
    import docx
except ImportError:
    print("Thư viện 'python-docx' chưa được cài đặt. Đang tiến hành cài đặt...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Đặt màu nền cho cell trong bảng"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)

def main():
    # Đường dẫn file
    current_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(current_dir, "CHATBOT_BUSINESS_FLOW.md")
    docx_path = os.path.join(current_dir, "CHATBOT_BUSINESS_FLOW.docx")

    if not os.path.exists(md_path):
        print(f"Không tìm thấy file: {md_path}")
        return

    print("Đang đọc nội dung file Markdown...")
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # Cấu hình lề trang (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Cấu hình Font chữ mặc định
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(51, 51, 51) # #333333

    print("Đang chuyển đổi nội dung sang Word...")
    
    in_mermaid = False
    
    for line_raw in lines:
        line = line_raw.strip()
        
        # Tự động tải ảnh và chèn sơ đồ Mermaid vào file Word
        if line.startswith("```mermaid"):
            in_mermaid = True
            mermaid_lines = []
            continue
        elif in_mermaid:
            if line.startswith("```"):
                in_mermaid = False
                try:
                    import base64
                    import urllib.request
                    
                    mermaid_code = "".join(mermaid_lines)
                    # Encode base64
                    encoded_code = base64.b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
                    url = f"https://mermaid.ink/img/{encoded_code}"
                    
                    temp_img_path = os.path.join(current_dir, "temp_mermaid.png")
                    print(f"Đang kết nối tới API để tải sơ đồ Mermaid...")
                    
                    req = urllib.request.Request(
                        url, 
                        headers={'User-Agent': 'Mozilla/5.0'}
                    )
                    with urllib.request.urlopen(req, timeout=10) as response:
                        with open(temp_img_path, 'wb') as out_file:
                            out_file.write(response.read())
                    
                    # Chèn hình ảnh vào Word
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_picture(temp_img_path, width=Inches(6.0))
                    
                    # Thêm chú thích cho ảnh
                    p_caption = doc.add_paragraph()
                    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_caption.paragraph_format.space_before = Pt(4)
                    p_caption.paragraph_format.space_after = Pt(12)
                    run_caption = p_caption.add_run("Hình 1: Sơ đồ luồng xử lý chi tiết của AI Chatbot")
                    run_caption.italic = True
                    run_caption.font.size = Pt(9.5)
                    run_caption.font.color.rgb = RGBColor(128, 128, 128)
                    
                    # Xóa ảnh tạm sau khi chèn xong
                    if os.path.exists(temp_img_path):
                        os.remove(temp_img_path)
                    print("Đã tải và chèn sơ đồ Mermaid thành công!")
                except Exception as e:
                    print(f"Lưu ý: Không tải được sơ đồ từ API (Có thể do thiết lập mạng): {e}")
                    p = doc.add_paragraph()
                    r = p.add_run("[SƠ ĐỒ LUỒNG XỬ LÝ - XEM TRONG FILE MARKDOWN CHATBOT_BUSINESS_FLOW.MD]")
                    r.italic = True
                    r.font.color.rgb = RGBColor(128, 128, 128)
            else:
                mermaid_lines.append(line_raw)
            continue
            
        # Tiêu đề H1
        if line_raw.startswith("# "):
            title_text = line_raw[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 51, 102) # Xanh Navy đậm
            
        # Tiêu đề H2
        elif line_raw.startswith("## "):
            h2_text = line_raw[3:].strip()
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.runs[0] if p.runs else p.add_run()
            run.text = h2_text
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 102, 204) # Xanh dương
            
        # Tiêu đề H3
        elif line_raw.startswith("### "):
            h3_text = line_raw[4:].strip()
            p = doc.add_heading(level=3)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.runs[0] if p.runs else p.add_run()
            run.text = h3_text
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(102, 102, 102) # Xám đậm
            
        # Bullet list hoặc list đánh số
        elif line.startswith("* ") or line.startswith("- ") or (line and line[0].isdigit() and ". " in line[:3]):
            is_bullet = line.startswith("* ") or line.startswith("- ")
            bullet_text = line[2:].strip() if is_bullet else line
            
            p = doc.add_paragraph(style='List Bullet' if is_bullet else 'List Number')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            
            # Xử lý in đậm cụm từ đầu dòng (ví dụ: * **Lọc cứng:** ...)
            parts = bullet_text.split("**")
            for i, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
                    
        # Dòng kẻ ngăn cách ---
        elif line == "---":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("____________________________________________________")
            run.font.color.rgb = RGBColor(200, 200, 200)
            
        # Đoạn văn thường
        elif line:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            
            parts = line.split("**")
            for i, part in enumerate(parts):
                if not part:
                    continue
                # Bỏ qua các định dạng markdown dạng code `brand`
                clean_part = part.replace("`", "")
                run = p.add_run(clean_part)
                if i % 2 == 1:
                    run.bold = True

    doc.save(docx_path)
    print(f"\n[Thành công] Đã xuất file Word ra địa chỉ:\n{docx_path}")

if __name__ == "__main__":
    main()
